"""
inserter.py — Image Inserter for Word
======================================
Run with: python inserter.py

Opens a small window. Pick images, type a title, choose columns,
click Insert — images are inserted exactly where your cursor is
in the open Word document.

Requires:
    pip install python-docx Pillow pywin32
"""

import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from PIL import Image, ImageTk
import os
from copy import deepcopy
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import win32com.client

# ─── CONFIG ──────────────────────────────────────────────────────────────────
PAGE_WIDTH_INCHES  = 6.5   # content width (8.5in page - 1in margins each side)
MIN_IMG_WIDTH_INCH = 1.5   # minimum image width — controls max auto columns
GUTTER_INCH        = 0.1   # gap between columns
CAPTION_FONT       = "Arial"
CAPTION_SIZE_PT    = 9
# ─────────────────────────────────────────────────────────────────────────────


def calc_columns(min_width_inch):
    cols = 1
    while True:
        next_cols = cols + 1
        col_w = (PAGE_WIDTH_INCHES - (next_cols - 1) * GUTTER_INCH) / next_cols
        if col_w < min_width_inch:
            break
        cols = next_cols
    return cols


def set_cell_border(cell):
    """Remove all visible borders from a table cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "nil")
        tcBorders.append(border)
    tcPr.append(tcBorders)


def insert_after(ref_elem, new_elem):
    """Insert new_elem immediately after ref_elem in the XML tree."""
    ref_elem.addnext(new_elem)


def insert_images_into_word(image_paths, title, num_cols):
    # ── Connect to Word ───────────────────────────────────────────────────────
    try:
        word_app = win32com.client.GetObject(Class="Word.Application")
    except Exception:
        raise Exception(
            "Could not connect to Word.\n\n"
            "Make sure Word is open with a document before clicking Insert."
        )

    if word_app.Documents.Count == 0:
        raise Exception("No documents are open in Word.\nPlease open a document first.")

    doc_path = word_app.ActiveDocument.FullName

    if doc_path.startswith("https://") or doc_path.startswith("http://"):
        raise Exception(
            "Your document is saved on OneDrive.\n\n"
            "Please go to File > Save As > Browse and save it locally "
            "(e.g. C:\\Users\\divya\\Documents) first."
        )

    # ── Record cursor position BEFORE closing ─────────────────────────────────
    cursor_pos = word_app.Selection.Range.Start

    # ── Save and close so python-docx can write to the file ───────────────────
    word_app.ActiveDocument.Save()
    word_app.ActiveDocument.Close()

    # ── Open with python-docx ─────────────────────────────────────────────────
    doc = Document(doc_path)

    # ── Find which paragraph the cursor was in ────────────────────────────────
    char_count = 0
    insert_idx = max(0, len(doc.paragraphs) - 1)  # fallback: last paragraph
    for i, para in enumerate(doc.paragraphs):
        para_len = len(para.text) + 1  # +1 for paragraph mark
        if char_count + para_len > cursor_pos:
            insert_idx = i
            break
        char_count += para_len

    ref_elem = doc.paragraphs[insert_idx]._element

    # ── Optionally insert a title paragraph at cursor ─────────────────────────
    if title.strip():
        # Build title para at end of doc, then move it to cursor position
        t_para = doc.add_paragraph()
        t_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = t_para.add_run(title.strip())
        run.bold      = True
        run.font.size = Pt(18)
        run.font.name = CAPTION_FONT
        t_elem = deepcopy(t_para._element)
        t_para._element.getparent().remove(t_para._element)
        insert_after(ref_elem, t_elem)
        ref_elem = t_elem  # table will go after title

    # ── Build image table ─────────────────────────────────────────────────────
    col_width  = (PAGE_WIDTH_INCHES - (num_cols - 1) * GUTTER_INCH) / num_cols
    table_cols = num_cols + (num_cols - 1)  # image cols + gutter cols between them
    table      = doc.add_table(rows=0, cols=table_cols)

    # Set column widths
    for i, col in enumerate(table.columns):
        is_gutter = (i % 2 == 1)
        width = Inches(GUTTER_INCH if is_gutter else col_width)
        for cell in col.cells:
            cell.width = width

    # Add rows of images
    for row_start in range(0, len(image_paths), num_cols):
        row = table.add_row()

        for col_idx in range(num_cols):
            table_col_idx = col_idx * 2  # account for gutter columns
            cell = row.cells[table_col_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            set_cell_border(cell)

            img_idx = row_start + col_idx
            if img_idx < len(image_paths):
                img_path = image_paths[img_idx]
                fname    = os.path.splitext(os.path.basename(img_path))[0]

                # Preserve aspect ratio
                with Image.open(img_path) as img:
                    orig_w, orig_h = img.size
                img_height = col_width * (orig_h / orig_w)

                # Image
                img_para = cell.paragraphs[0]
                img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img_para.add_run().add_picture(
                    img_path,
                    width=Inches(col_width),
                    height=Inches(img_height)
                )

                # Caption with bottom border
                cap_para = cell.add_paragraph()
                cap_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

                pPr    = cap_para._p.get_or_add_pPr()
                pBdr   = OxmlElement("w:pBdr")
                bottom = OxmlElement("w:bottom")
                bottom.set(qn("w:val"), "single")
                bottom.set(qn("w:sz"), "4")
                bottom.set(qn("w:space"), "1")
                bottom.set(qn("w:color"), "CCCCCC")
                pBdr.append(bottom)
                pPr.append(pBdr)

                label = cap_para.add_run("Caption: ")
                label.bold           = True
                label.font.size      = Pt(CAPTION_SIZE_PT)
                label.font.name      = CAPTION_FONT
                label.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

                name = cap_para.add_run(f"[{fname}]")
                name.italic          = True
                name.font.size       = Pt(CAPTION_SIZE_PT)
                name.font.name       = CAPTION_FONT
                name.font.color.rgb  = RGBColor(0xAA, 0xAA, 0xAA)

            else:
                # Empty padding cell for odd number of images
                set_cell_border(cell)

            # Remove gutter cell border
            if col_idx < num_cols - 1:
                set_cell_border(row.cells[table_col_idx + 1])

    # ── Move table to cursor position ─────────────────────────────────────────
    tbl_elem = deepcopy(table._element)
    table._element.getparent().remove(table._element)
    insert_after(ref_elem, tbl_elem)

    # Add empty paragraph after table for spacing
    sp = OxmlElement("w:p")
    insert_after(tbl_elem, sp)

    # ── Save and reopen ───────────────────────────────────────────────────────
    doc.save(doc_path)
    word_app.Documents.Open(doc_path)
    word_app.Visible = True
    word_app.Activate()


# ═══════════════════════════════════════════════════════════════════════════════
# GUI
# ═══════════════════════════════════════════════════════════════════════════════

class ImageInserterApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Image Inserter for Word")
        self.root.resizable(True, True)
        self.root.configure(bg="#f0f0f0")
        self.root.minsize(420, 680)
        self.root.geometry("440x720")

        self.image_paths = []
        self.thumb_refs  = []
        self._build_ui()

    def _build_ui(self):
        # Header
        tk.Label(self.root, text="📷 Image Inserter for Word",
                 font=("Segoe UI", 13, "bold"), bg="#f0f0f0").pack(pady=(16, 2))
        tk.Label(self.root, text="Images insert at your cursor position in Word.",
                 font=("Segoe UI", 9), fg="#666", bg="#f0f0f0").pack(pady=(0, 8))
        ttk.Separator(self.root, orient="horizontal").pack(fill="x", padx=14)

        # Title field
        tk.Label(self.root, text="Document Title (optional)",
                 font=("Segoe UI", 9, "bold"), bg="#f0f0f0", anchor="w").pack(fill="x", padx=14, pady=(10,2))
        self.title_var = tk.StringVar()
        tk.Entry(self.root, textvariable=self.title_var, font=("Segoe UI", 10),
                 relief="solid", bd=1).pack(fill="x", padx=14)

        # Columns
        tk.Label(self.root, text="Columns per row",
                 font=("Segoe UI", 9, "bold"), bg="#f0f0f0", anchor="w").pack(fill="x", padx=14, pady=(10,2))
        col_frame = tk.Frame(self.root, bg="#f0f0f0")
        col_frame.pack(fill="x", padx=14)
        self.col_var = tk.StringVar(value="Auto")
        for val in ["Auto", "2", "3", "4"]:
            tk.Radiobutton(col_frame, text=val, variable=self.col_var, value=val,
                           font=("Segoe UI", 9), bg="#f0f0f0").pack(side="left", padx=6)

        # Image buttons
        tk.Label(self.root, text="Images",
                 font=("Segoe UI", 9, "bold"), bg="#f0f0f0", anchor="w").pack(fill="x", padx=14, pady=(10,2))
        btn_frame = tk.Frame(self.root, bg="#f0f0f0")
        btn_frame.pack(fill="x", padx=14)
        tk.Button(btn_frame, text="+ Add Images", command=self.add_images,
                  font=("Segoe UI", 9), bg="#2b7cd3", fg="white",
                  relief="flat", padx=10, pady=4).pack(side="left")
        tk.Button(btn_frame, text="Clear All", command=self.clear_images,
                  font=("Segoe UI", 9), bg="#e0e0e0",
                  relief="flat", padx=10, pady=4).pack(side="left", padx=8)

        self.count_label = tk.Label(self.root, text="No images selected",
                                    font=("Segoe UI", 9), fg="#888", bg="#f0f0f0")
        self.count_label.pack(anchor="w", padx=14, pady=(4,2))

        # Bottom — Insert button + status (pinned to bottom so always visible)
        bottom_frame = tk.Frame(self.root, bg="#f0f0f0")
        bottom_frame.pack(side="bottom", fill="x", padx=14, pady=(6,12))

        self.status_var = tk.StringVar()
        self.status_label = tk.Label(bottom_frame, textvariable=self.status_var,
                                     font=("Segoe UI", 9), fg="#2d7a3a", bg="#f0f0f0",
                                     wraplength=390)
        self.status_label.pack(pady=(0,6))

        tk.Button(bottom_frame, text="Insert into Word Document",
                  command=self.do_insert,
                  font=("Segoe UI", 11, "bold"),
                  bg="#2b7cd3", fg="white",
                  relief="flat", pady=10).pack(fill="x")

        # Scrollable image preview (middle, fills remaining space)
        preview_frame = tk.Frame(self.root, bg="white", relief="solid", bd=1)
        preview_frame.pack(fill="both", expand=True, padx=14, pady=(0,6))

        self.canvas = tk.Canvas(preview_frame, bg="white", highlightthickness=0)
        scrollbar   = ttk.Scrollbar(preview_frame, orient="vertical", command=self.canvas.yview)
        self.canvas.configure(yscrollcommand=scrollbar.set)
        scrollbar.pack(side="right", fill="y")
        self.canvas.pack(side="left", fill="both", expand=True)

        self.thumb_frame = tk.Frame(self.canvas, bg="white")
        self.canvas.create_window((0,0), window=self.thumb_frame, anchor="nw")
        self.thumb_frame.bind("<Configure>", lambda e: self.canvas.configure(
            scrollregion=self.canvas.bbox("all")))

    def add_images(self):
        files = filedialog.askopenfilenames(
            title="Select Images",
            filetypes=[("Image files", "*.jpg *.jpeg *.png *.bmp *.gif *.webp"), ("All files", "*.*")]
        )
        for f in files:
            if f not in self.image_paths:
                self.image_paths.append(f)
        self.refresh_thumbnails()

    def clear_images(self):
        self.image_paths = []
        self.refresh_thumbnails()
        self.status_var.set("")

    def refresh_thumbnails(self):
        for widget in self.thumb_frame.winfo_children():
            widget.destroy()
        self.thumb_refs = []

        count = len(self.image_paths)
        self.count_label.config(
            text=f"{count} image{'s' if count != 1 else ''} selected" if count > 0 else "No images selected"
        )

        for i, path in enumerate(self.image_paths):
            row_frame = tk.Frame(self.thumb_frame, bg="white")
            row_frame.pack(fill="x", padx=4, pady=2)

            try:
                img = Image.open(path)
                img.thumbnail((40, 40))
                photo = ImageTk.PhotoImage(img)
                self.thumb_refs.append(photo)
                tk.Label(row_frame, image=photo, bg="white").pack(side="left", padx=4)
            except Exception:
                tk.Label(row_frame, text="🖼", font=("Segoe UI", 16), bg="white").pack(side="left", padx=4)

            tk.Label(row_frame, text=os.path.basename(path), font=("Segoe UI", 9),
                     bg="white", anchor="w", fg="#333").pack(side="left", fill="x", expand=True)

            tk.Button(row_frame, text="×", command=lambda i=i: self.remove_image(i),
                      font=("Segoe UI", 10), fg="#999", bg="white",
                      relief="flat", bd=0).pack(side="right", padx=4)

    def remove_image(self, idx):
        if 0 <= idx < len(self.image_paths):
            self.image_paths.pop(idx)
            self.refresh_thumbnails()

    def do_insert(self):
        if not self.image_paths:
            messagebox.showwarning("No images", "Please add at least one image first.")
            return

        col_str  = self.col_var.get()
        num_cols = calc_columns(MIN_IMG_WIDTH_INCH) if col_str == "Auto" else int(col_str)

        self.status_var.set("Inserting… please wait.")
        self.status_label.config(fg="#1a56a0")
        self.root.update()

        try:
            insert_images_into_word(self.image_paths, self.title_var.get(), num_cols)
            n = len(self.image_paths)
            self.status_var.set(f"✅ {n} image{'s' if n != 1 else ''} inserted at cursor!")
            self.status_label.config(fg="#2d7a3a")
        except Exception as e:
            self.status_var.set(f"❌ {e}")
            self.status_label.config(fg="#b71c1c")


if __name__ == "__main__":
    root = tk.Tk()
    ImageInserterApp(root)
    root.mainloop()
